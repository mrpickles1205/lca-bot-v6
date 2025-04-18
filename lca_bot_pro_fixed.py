
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import datetime

def generate_lca_report(product):
    df = pd.DataFrame({
        'Life Cycle Stage': ['Materials', 'Manufacturing', 'Use Phase', 'End-of-Life'],
        'Energy Use (MJ)': [random.uniform(80, 120), random.uniform(50, 100), random.uniform(10, 20), random.uniform(15, 30)],
        'GHG Emissions (kg CO2-eq)': [random.uniform(5, 10), random.uniform(8, 12), random.uniform(1, 3), random.uniform(2, 4)],
        'Water Use (L)': [random.uniform(20, 40), random.uniform(10, 30), random.uniform(1, 5), random.uniform(5, 15)]
    })

    chart_paths = []
    for metric in df.columns[1:]:
        fig, ax = plt.subplots()
        ax.bar(df['Life Cycle Stage'], df[metric])
        ax.set_title(f'{metric} by Stage')
        chart_file = f"{metric.replace(' ', '_')}_chart.png"
        fig.savefig(chart_file)
        chart_paths.append(chart_file)
        plt.close(fig)

    doc = Document()
    doc.add_heading(f'LCA Report for: {product}', 0)
    doc.add_paragraph(f'Date: {datetime.date.today()}')
    doc.add_paragraph("Confidential – For Internal Use Only").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    toc_sections = [
        "Executive Summary", "1. Introduction", "2. Goal and Scope", "3. Functional Unit", "4. System Boundary",
        "5. Inventory Analysis", "6. Life Cycle Impact Assessment (LCIA)", "7. Interpretation",
        "8. Assumptions and Limitations", "9. Recommendations", "Appendix A: Glossary", "Appendix B: References"
    ]
    for section in toc_sections:
        doc.add_paragraph(section)
    doc.add_page_break()

    sections = {
        "Executive Summary": """This report presents an ISO-compliant life cycle assessment of one electric toothbrush.
The goal is to assess its environmental impacts from cradle to grave. Key impact categories include energy use,
greenhouse gas emissions, and water consumption. Findings suggest that the manufacturing and end-of-life stages are
most impactful. Improvement strategies are discussed.""",

        "1. Introduction": """This Life Cycle Assessment (LCA) has been conducted in accordance with ISO 14040 and 14044.
The aim is to assess and document the environmental footprint of an electric toothbrush product. This document provides
structured analysis from material sourcing through disposal.""",

        "2. Goal and Scope": """The LCA's objective is to evaluate the cradle-to-grave environmental performance of the product.
Scope includes raw material extraction, manufacturing, packaging, distribution, use phase, and end-of-life treatment.""",

        "3. Functional Unit": """The functional unit is defined as one electric toothbrush used over a 3-year lifetime,
assuming daily use and weekly recharging. All inputs and outputs are normalized to this unit.""",

        "4. System Boundary": """The assessment uses a cradle-to-grave system boundary, including:
- Material extraction and processing
- Manufacturing and assembly
- Use phase (electricity consumption)
- Transportation and packaging
- End-of-life scenarios including recycling and landfill""",

        "5. Inventory Analysis": """The inventory includes material composition (plastic, metal, battery), manufacturing steps,
transport energy, electricity use, and end-of-life disposal paths. Data is based on literature and approximations.""",

        "6. Life Cycle Impact Assessment (LCIA)": """LCIA converts inventory data into potential environmental impacts.
Categories assessed:
- Energy demand (MJ)
- Global warming potential (kg CO2-eq)
- Water use (liters)
Methods used: ReCiPe 2016 Midpoint (H) and IPCC 2021 for GWP.

Impact results are shown below.""" ,

        "7. Interpretation": """The manufacturing stage contributes most to energy demand and emissions, largely due to battery and
circuit production. End-of-life impact varies by disposal method. Use-phase energy is relatively minor but cumulative.""" ,

        "8. Assumptions and Limitations": """Assumptions include average electricity grid mix, transport distances, and usage frequency.
Limitations include lack of real supplier-specific data and end-of-life variability. Cutoff criteria: flows <5% impact excluded.""" ,

        "9. Recommendations": """Design for disassembly, replaceable batteries, and recycled materials could significantly reduce impacts.
Consumer education on e-waste recycling and product durability improvements are also recommended.""" ,

        "Appendix A: Glossary": """LCA: Life Cycle Assessment
GWP: Global Warming Potential
MJ: Megajoules
CO2-eq: Carbon dioxide equivalent""" ,

        "Appendix B: References": """1. ISO 14040:2006
2. ISO 14044:2006
3. ReCiPe 2016 Method
4. IPCC AR6 Methodology
5. Ecoinvent Database"""
    }

    for title, content in sections.items():
        doc.add_heading(title, level=1 if not title.startswith("Appendix") else 2)
        for paragraph in content.split('\n'):
            doc.add_paragraph(paragraph)
        if title == "6. Life Cycle Impact Assessment (LCIA)":
            for chart in chart_paths:
                doc.add_picture(chart, width=Inches(5.5))
        doc.add_page_break()

    doc.add_heading("Detailed LCI Table", level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(round(val, 2)) if isinstance(val, (int, float)) else str(val)

    filename = f"LCA_Report_Pro_{product.replace(' ', '_')}.docx"
    doc.save(filename)
    print(f"Report saved as {filename}")

if __name__ == "__main__":
    generate_lca_report("Electric Toothbrush")
